# This script creates a text report for BRCA-analyzer

enLang={'heading0':'BRCA-analyzer Report',
        'heading1':'1. Read Statistics',
        'TotalNum':'Total Number of Reads: ',
        'IndexedReads':'Number of reads with determined index (barcode): ',
        'LowIndexedReads1':'WARNING! Too low percent of indexed reads (<50%)!',
        'LowIndexedReads2':'WARNING! Low percent of indexed reads (<80%)!',
        'EmptyReads':'Number of reads for empty indexes: ',
        'HighEmpty':'WARNING! Too high number of reads for empty indexes!',
        'MedianTrimmedReads':'Median percent of reads with properly removed primer sequences: ',
        'LowMedianTrimmedReads1':'WARNING! Too low percent of reads with properly removed primer sequences (<50%)!',
        'LowMedianTrimmedReads2':'WARNING! Low percent of reads with properly removed primer sequences (<80%)!',
        'WarningFor':'WARNING! For ',
        'TooLowTrimmedReadsWarn1':' sample(s) there is too low percent of reads with properly removed primer sequences (<50%)!',
        'TooLowTrimmedReadsWarn2':' sample(s) there is low percent of reads with properly removed primer sequences (<80%, but >=50%)!',
        'heading2':'2. Coverage Statistics',
        'MedAmpliconCov':'Median amplicon coverage: ',
        'MedLowCovedAmpls':'Median number of weakly covered amplicons (<30 reads): ',
        'LowMedAmpliconCov':'WARNING! Too low median amplicon coverage (<100 reads)!',
        'HighMedLowCovedAmpls':'WARNING! For the most of samples too many amplicons (>10) has low coverage (<30)!',
        'HighNumLowCovedAmpls':' samples there is more than 10 amplicons covered with low number of reads (<30)!',
        'ModNumLowCovedAmpls':' samples there is from 5 to 10 amplicons covered with low number of reads (<30)!',
        'LowAmplCovs':' amplicons median coverage is too low (<100 reads)!',
        'TotalUncoveredAmpls':'Total number of amplicon(s) with low coverage (<30 reads): ',
        'HighLowCovedAmplsPercWarn':'WARNING: Too high total percent of amplicons with low coverage (<30 reads)!',
        'heading3':'3. Found Variants',
        'PatientID':'Patient ID','Gene':'Gene','CDS':'CDS','Protein':'Protein','Quality':'Quality','ClinSign':'Clin. Sign.','ClassReasons':'Reasons for Class.',
        'Comment':'Comment','QualHigh':'HIGH','QualMod':'MODER.','QualLow':'LOW','Contamination':'Contam.','InSilicoTools':'predicted in silico',
        'VarPercent':'%','Designations':'Designations:',
        'CDS_Def':'changes in the coding sequence of gene',
        'Protein_Def':'changes in the protein aminoacid sequence',
        'VarPercent_Def':'percent of reads that contain alternative allele',
        'ClinSign_Def':'clinical significance of the variant (5-Pathogenic, 4-Likely pathogenic, 3-Uncertain significance)',
        'ClassReasons_Def':'reasons for classification this variant as defined',
        'Contamination_Def':'possible contamination (between samples or sequencing runs)',
        'Warnings':'Warnings:','VarForCheckWarn1':'Variants with numbers ','VarForCheckWarn2':' need additional check with Sanger sequencing!',
        'VarsContamWarn2':' may be due to contamination!',
        'PatsWithSeveralMutsWarn1':'For sample(s) ','PatsWithSeveralMutsWarn2':' there are more than 1 variants found! This is a rare case, check these samples.',
        'from':'from','to':'to','#':'#'}
ruLang={'heading0':'Отчет обработки данных NGS с помощью BRCA-analyzer',
        'heading1':'1. Статистика по качеству и количеству прочтениям',
        'TotalNum':'Всего прочтений: ',
        'IndexedReads':'Число прочтений с определенным индексом (баркодом): ',
        'LowIndexedReads1':'ВНИМАНИЕ! Значительно снижено число прочтений с определенным индексом!',
        'LowIndexedReads2':'ВНИМАНИЕ! Снижено число прочтений с определенным индексом!',
        'EmptyReads':'Число прочтений для пустых индексов: ',
        'HighEmpty':'ВНИМАНИЕ! Много прочтений для пустых индексов!',
        'MedianTrimmedReads':'Медианное число прочтений с идентифицированными последовательностями праймеров: ',
        'LowMedianTrimmedReads1':'ВНИМАНИЕ! Значительно снижен процент прочтений с правильно идентифицированными последовательностями праймеров (<50%)!',
        'LowMedianTrimmedReads2':'ВНИМАНИЕ! Низкий процент прочтений с правильно идентифицированными последовательностями праймеров (<80%)!',
        'WarningFor':'ВНИМАНИЕ! Для ',
        'TooLowTrimmedReadsWarn1':' образцов(-а) значительно снижен процент прочтений с правильно идентифицированными последовательностями праймеров (<50%)!',
        'TooLowTrimmedReadsWarn2':' образцов(-а) снижен процент прочтений с правильно идентифицированными последовательностями праймеров (<80%, но >=50%)!',
        'heading2':'2. Статистика по покрытию целевых последовательностей',
        'MedAmpliconCov':'Медианное покрытие ампликонов: ',
        'MedLowCovedAmpls':'Медианное число ампликонов с покрытием менее 30 прочтений: ',
        'LowMedAmpliconCov':'ВНИМАНИЕ! Значительно снижено медианное покрытие ампликонов (<100)!',
        'HighMedLowCovedAmpls':'ВНИМАНИЕ! Для большинства образцов большое число ампликонов (>10) имеют низкое покрытие (<30)!',
        'HighNumLowCovedAmpls':' образцов(-а) более 10 ампликонов покрыты менее, чем 30 прочтениями!',
        'ModNumLowCovedAmpls':' образцов(-а) от 5 до 10 ампликонов покрыты менее, чем 30 прочтениями!',
        'LowAmplCovs':' ампликонов(-а) медианное покрытие снижено (менее 100 прочтений)!',
        'TotalUncoveredAmpls':'Всего ампликонов с низким покрытием (<30 прочтений): ',
        'HighLowCovedAmplsPercWarn':'ВНИМАНИЕ! Значительный процент ампликонов имеют покрытие менее 30 прочтений!',
        'heading3':'3. Найденные варианты',
        'PatientID':'ID пациента','Gene':'Ген','CDS':'Кодир. посл.','Protein':'Аминокис. посл.','Quality':'Качество','ClinSign':'Клин. знач.','ClassReasons':'Признаки клин. знач.',
        'Comment':'Коммент.','QualHigh':'ВЫС.','QualMod':'СРЕД.','QualLow':'НИЗ.','Contamination':'Контам.','InSilicoTools':'предсказано in silico',
        'VarPercent':'%','Designations':'Обозначения:',
        'CDS_Def':'изменение в кодирующей последовательности гена',
        'Protein_Def':'изменение к аминокислотной последовательности кодируемого геном белка',
        'VarPercent_Def':'процент прочтений, содержащих альтернативный аллель (отличный от референсного)',
        'ClinSign_Def':'клиническая значимость варианта (5-клинически значимый, 4-вероятно клинически значимый, 3-неизвестное клиническое значение)',
        'ClassReasons_Def':'признаки, по которым данное клинеское значение было определено',
        'Contamination_Def':'возможная контаминация (между образцами или между запусками секвенирования',
        'Warnings':'Предупреждения:','VarForCheckWarn1':'Для вариантов с номерами ','VarForCheckWarn2':' необходима дополнительная проверка секвенированием по Сэнгеру!',
        'VarsContamWarn2':' могут быть следствием контаминации!',
        'PatsWithSeveralMutsWarn1':'Для образцов(-а) ','PatsWithSeveralMutsWarn2':' найдено более 1 варианта! Это редкий случай, проверьте данные образцы.',
        'from':'от','to':'до','#':'№'}

import xlrd
import docx
import argparse
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt,Mm,Inches,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT,WD_SECTION
import traceback
from docx.oxml.shared import OxmlElement, qn
import statistics as stat
import numpy

def changeStyles(doc):
    # Change zoom to 100%
    doc.settings.element[0].set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}percent','100')
    # Change margins
    for i,sec in enumerate(doc.sections):
        if i==0:
            sec.top_margin=Mm(20)
            sec.bottom_margin=Mm(20)
            sec.left_margin=Mm(30)
            sec.right_margin=Mm(15)
        else:
            sec.top_margin=Mm(30)
            sec.bottom_margin=Mm(15)
            sec.left_margin=Mm(20)
            sec.right_margin=Mm(20)
    # Title
    style=doc.styles['Title']
    style.element[7].getchildren()[0].getchildren()[0].set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color','000000')
    font=style.font
    font.name='Times New Roman'
    font.size=Pt(16)
    font.italic=False
    font.bold=True
    font.underline=False
    font.color.rgb=RGBColor(0,0,0)
    style.paragraph_format.space_after=Pt(12)
    style.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    # Heading 1
    style=doc.styles['Heading 1']
    font=style.font
    font.name='Times New Roman'
    font.size=Pt(14)
    font.italic=False
    font.bold=True
    font.underline=False
    font.color.rgb=RGBColor(0,0,0)
    style.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.space_after=Pt(12)
    # Heading 2
    style=doc.styles['Heading 2']
    font=style.font
    font.name='Times New Roman'
    font.size=Pt(14)
    font.italic=False
    font.bold=True
    font.underline=False
    font.color.rgb=RGBColor(0,0,0)
    style.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.space_after=Pt(6)
    # Normal
    style=doc.styles['Normal']
    font=style.font
    font.name='Times New Roman'
    font.size=Pt(12)
    font.italic=False
    font.bold=False
    font.underline=False
    font.color.rgb=RGBColor(0,0,0)
    style.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after=Pt(0)
    style.paragraph_format.space_before=Pt(0)

def set_cell_vertical_alignment(cell, align="center"): 
    try:   
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcValign = OxmlElement('w:vAlign')  
        tcValign.set(qn('w:val'), align)  
        tcPr.append(tcValign)
        return True 
    except:
        traceback.print_exc()             
        return False

par=argparse.ArgumentParser(description='This script creates report about BRCA-analyzer results')
par.add_argument('--read-stat-file','-read',dest='readStatFile',type=str,help='file with statistics of reads',required=True)
par.add_argument('--cov-stat-file','-cov',dest='covStatFile',type=str,help='file with statistics of coverage',required=True)
par.add_argument('--result-file','-res',dest='resFile',type=str,help='file with BRCA-analyzer results of variant calling',required=True)
par.add_argument('--uniformity-figure','-fig',dest='figFile',type=str,help='figure with uniformity of coverage',required=True)
par.add_argument('--output-file','-out',dest='outFile',type=str,help='file for output',required=True)
par.add_argument('--language','-lang',dest='lang',type=str,help='language of report (russian or english). Default: english',default='english')
args=par.parse_args()

langs=['russian','english']
if args.lang not in langs:
    print('#'*10,'\nWARNING! Chosen language is not accepted. Use default english...')
if args.lang=='russian': lang=ruLang
else: lang=enLang
file=open(args.readStatFile)
indexed=[]
allTotalReads=[]
allNotEmpty=[]
empty=[]
trimmedReadsPerc=[]
tooLowTrimmedReads=0
lowTrimmedReads=0
trimPrimerStat=False
for string in file:
    if 'Patient_Num' in string:
        if 'Properly_Trimmed' in string:
            trimPrimerStat=True
        continue
    if string=='' or string=='\n': break
    cols=string.replace('\n','').split('\t')
    allTotalReads.append(int(cols[4]))
    if 'Undetermined' in cols[0]:
        continue
    indexed.append(int(cols[4]))
    if 'empty' in cols[1]:
        empty.append(int(cols[4]))
    else:
        allNotEmpty.append(int(cols[4]))
        if trimPrimerStat:
            trimmedReadsPerc.append(float(cols[6]))
            if float(cols[6])<0.50:
                tooLowTrimmedReads+=1
            elif 0.50<=float(cols[6])<0.80:
                lowTrimmedReads+=1
file.close()
# Indexed reads
percIndexed=round(sum(indexed)*100/sum(allTotalReads),1)
lowIndexedReads=''; lowIndexedReadsColor=RGBColor(255,255,255)
if percIndexed<50: lowIndexedReads=lang['LowIndexedReads1']; lowIndexedReadsColor=RGBColor(255,0,0)
elif 50<=percIndexed<80: lowIndexedReads=lang['LowIndexedReads2']; lowIndexedReadsColor=RGBColor(255,78,0)
# Empty indexes
percEmpty=round(sum(empty)*100/sum(allTotalReads),2)
highEmpty=''; highEmptyColor=RGBColor(255,255,255)
if max(empty)>min(allNotEmpty):
    highEmpty=lang['HighEmpty']; highEmptyColor=RGBColor(255,0,0)
# Trimming reads stat
medianTrimmedRead=round(stat.median(trimmedReadsPerc)*100,1)
minTrimmedRead=round(min(trimmedReadsPerc)*100,1); maxTrimmedRead=round(max(trimmedReadsPerc)*100,1)
lowMedianTrimmedReads=''; lowMedianTrimmedReadsColor=RGBColor(255,255,255)
if medianTrimmedRead<50: lowMedianTrimmedReads=lang['LowMedianTrimmedReads1']; lowMedianTrimmedReadsColor=RGBColor(255,0,0)
elif 50<=percIndexed<80: lowMedianTrimmedReads=lang['LowMedianTrimmedReads2']; lowMedianTrimmedReadsColor=RGBColor(255,78,0)
tooLowTrimmedReadsWarn=''; tooLowTrimmedReadsWarnColor=RGBColor(255,255,255)
if tooLowTrimmedReads>0:
    tooLowTrimmedReadsWarn=lang['WarningFor']+str(tooLowTrimmedReads)+lang['TooLowTrimmedReadsWarn1']
    tooLowTrimmedReadsWarnColor=RGBColor(255,0,0)
lowTrimmedReadsWarn=''; lowTrimmedReadsWarnColor=RGBColor(255,255,255)
if lowTrimmedReads>0:
    lowTrimmedReadsWarn=lang['WarningFor']+str(lowTrimmedReads)+lang['TooLowTrimmedReadsWarn2']
    lowTrimmedReadsWarnColor=RGBColor(255,78,0)
# Coverage
file=open(args.covStatFile)
medCovs=[]
lowCovedAmpls=[]
totalAmplsNum=0
highNumLowCovedAmpls=[]
modNumLowCovedAmpls=[]
vals=[]
for string in file:
    if 'Patient#' in string: continue
    cols=string.replace('\n','').split('\t')
    if 'empty' in cols[1] or 'notemplatecontrol' in cols[1]: continue
    if string=='' or string=='\n': break
    medCovs.append(float(cols[3]))
    lowCovedAmpls.append(int(cols[4]))
    if int(cols[4])>10:
        highNumLowCovedAmpls.append(int(cols[4]))
    elif 5<=int(cols[4])<=10:
        modNumLowCovedAmpls.append(int(cols[4]))
    vals.append(list(map(float,cols[5:])))
    totalAmplsNum+=len(cols[5:])
lowCovedAmplsPerc=round(sum(lowCovedAmpls)*100/totalAmplsNum,1)
valsArray=numpy.array(vals)
amplCovs=[]
lowAmplCovs=[]
for j in range(valsArray.shape[1]):
    amplCovs.append(stat.median(valsArray[:,j]))
    if stat.median(valsArray[:,j])<100:
        lowAmplCovs.append(j+1)
medAmpliconCov=round(stat.median(medCovs),1); minAmpliconCov=round(min(medCovs),1); maxAmpliconCov=round(max(medCovs),1)
if medAmpliconCov<100:
    lowMedAmpliconCov=lang['LowMedAmpliconCov']; lowMedAmpliconCovColor=RGBColor(255,0,0)
medLowCovedAmpls=round(stat.median(lowCovedAmpls),0); minLowCovedAmpls=round(min(lowCovedAmpls),0); maxLowCovedAmpls=round(max(lowCovedAmpls),0)
if medLowCovedAmpls>10:
    highMedLowCovedAmpls=lang['HighMedLowCovedAmpls']; highMedLowCovedAmplsColor=RGBColor(255,0,0)
if len(highNumLowCovedAmpls)>0:
    highNumLowCovedAmplsWarn=lang['WarningFor']+str(len(highNumLowCovedAmpls))+lang['HighNumLowCovedAmpls']
    highNumLowCovedAmplsWarnColor=RGBColor(255,0,0)
if len(modNumLowCovedAmpls)>0:
    modNumLowCovedAmplsWarn=lang['WarningFor']+str(len(modNumLowCovedAmpls))+lang['ModNumLowCovedAmpls']
    modNumLowCovedAmplsWarnColor=RGBColor(255,78,0)
if len(lowAmplCovs)>0:
    lowAmplCovsWarn=lang['WarningFor']+str(len(lowAmplCovs))+lang['LowAmplCovs']
    lowAmplCovsWarnColor=RGBColor(255,0,0)
if lowCovedAmplsPerc>5:
    highLowCovedAmplsPercWarn=lang['HighLowCovedAmplsPercWarn']; highLowCovedAmplsPercWarnColor=RGBColor(255,0,0)
# Creating document
doc=docx.Document()
section = doc.sections[0]
section.orientation=WD_ORIENT.PORTRAIT
section.page_height = Mm(297)
section.page_width = Mm(210)
# Write title of document
head=doc.add_heading(lang['heading0'],0)
# Write first part (Reads statistics)
head1=doc.add_heading(lang['heading1'],1)
p=doc.add_paragraph(lang['TotalNum']+str(sum(allTotalReads)),style='List Bullet')
p=doc.add_paragraph(lang['IndexedReads']+str(sum(indexed))+' ('+str(percIndexed)+'%) ',style='List Bullet')
r=p.add_run(lowIndexedReads); r.font.color.rgb=lowIndexedReadsColor; r.font.bold=True
if len(empty)>0:
    p=doc.add_paragraph(lang['EmptyReads']+str(sum(empty))+' ('+str(percEmpty)+'%) ',style='List Bullet')
    r=p.add_run(highEmpty); r.font.color.rgb=highEmptyColor; r.font.bold=True
if trimPrimerStat:
    p=doc.add_paragraph(lang['MedianTrimmedReads']+str(medianTrimmedRead)+'% ('+lang['from']+' '+str(minTrimmedRead)+'% '+lang['to']+' '+str(maxTrimmedRead)+'%) ',style='List Bullet')
    r=p.add_run(lowMedianTrimmedReads); r.font.color.rgb=lowMedianTrimmedReadsColor; r.font.bold=True
    if tooLowTrimmedReads>0:
        p=doc.add_paragraph('',style='List Bullet'); r=p.add_run(tooLowTrimmedReadsWarn)
        r.font.color.rgb=tooLowTrimmedReadsWarnColor; r.font.bold=True
    if lowTrimmedReads>0:
        p=doc.add_paragraph('',style='List Bullet'); r=p.add_run(lowTrimmedReadsWarn)
        r.font.color.rgb=lowTrimmedReadsWarnColor; r.font.bold=True
# Write second part (Coverage statistics)
head2=doc.add_heading(lang['heading2'],1)
p=doc.add_paragraph(lang['MedAmpliconCov']+str(medAmpliconCov)+' ('+lang['from']+' '+str(minAmpliconCov)+' '+lang['to']+' '+str(maxAmpliconCov)+')',style='List Bullet')
if medAmpliconCov<100:
    r=p.add_run(lowMedAmpliconCov); r.font.color.rgb=lowMedAmpliconCovColor; r.font.bold=True
p=doc.add_paragraph(lang['MedLowCovedAmpls']+str(medLowCovedAmpls)+' ('+lang['from']+' '+str(minLowCovedAmpls)+' '+lang['to']+' '+str(maxLowCovedAmpls)+')',style='List Bullet')
if medLowCovedAmpls>10:
    r=p.add_run(highMedLowCovedAmpls); r.font.color.rgb=highMedLowCovedAmplsColor; r.font.bold=True
p=doc.add_paragraph(lang['TotalUncoveredAmpls']+str(sum(lowCovedAmpls))+' ('+str(lowCovedAmplsPerc)+'%) ',style='List Bullet')
if lowCovedAmplsPerc>5:
    r=p.add_run(highLowCovedAmplsPercWarn); r.font.color.rgb=highLowCovedAmplsPercWarnColor; r.font.bold=True
if len(highNumLowCovedAmpls)>0:
    p=doc.add_paragraph('',style='List Bullet'); r=p.add_run(highNumLowCovedAmplsWarn)
    r.font.color.rgb=highNumLowCovedAmplsWarnColor; r.font.bold=True
if len(modNumLowCovedAmpls)>0:
    p=doc.add_paragraph('',style='List Bullet'); r=p.add_run(modNumLowCovedAmplsWarn)
    r.font.color.rgb=modNumLowCovedAmplsWarnColor; r.font.bold=True
if len(lowAmplCovs)>0:
    p=doc.add_paragraph('',style='List Bullet'); r=p.add_run(lowAmplCovsWarn)
    r.font.color.rgb=lowAmplCovsWarnColor; r.font.bold=True
doc.add_picture(args.figFile,width=Mm(160))
# Write third part (Found variants)
section=doc.add_section(WD_SECTION.NEW_PAGE)
section.orientation=WD_ORIENT.LANDSCAPE
section.page_height=Mm(210)
section.page_width=Mm(297)
head3=doc.add_heading(lang['heading3'],1)
table=doc.add_table(rows=1,cols=10)
table.style = 'Table Grid'
hdr_cells=table.rows[0].cells
colNames=[lang['#'],lang['PatientID'],lang['Gene'],lang['CDS'],lang['Protein'],lang['Quality'],lang['VarPercent'],lang['ClinSign'],lang['ClassReasons'],lang['Comment']]
colWidths=[0.5,1.6,0.8,1.4,1.2,1,0.7,0.7,1.2,1]
for i,colName in enumerate(colNames):
    table.columns[i].width=Inches(colWidths[i])
    hdr_cells[i].text=colName
    hdr_cells[i].width=Inches(colWidths[i])
    run=hdr_cells[i].paragraphs[0].runs[0]
    run.bold=True; hdr_cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_cell_vertical_alignment(hdr_cells[i])
# Read result table
wb=xlrd.open_workbook(args.resFile)
varsForCheck=[]
varsContam=[]
patVarNums={}
rowNum=0
# Pathogenic
ws=wb.sheet_by_name('Pathogenic')
muts=[]
for i in range(ws.nrows)[1:]:
    row=ws.row_values(i)
    muts.append(row[13])
for i in range(ws.nrows)[1:]:
    if row[1] not in patVarNums.keys():
        patVarNums[row[1]]=1
    else:
        patVarNums[row[1]]+=1
    row=ws.row_values(i)
    comment=''; commentColor=RGBColor(0,0,0)
    if float(row[8])>=1200 and float(row[18])>=0.14: qual=lang['QualHigh']; qualColor=RGBColor(0,100,0)
    elif float(row[8])<500 and float(row[18])<0.14:
        varsForCheck.append(str(i+rowNum))
        qual=lang['QualLow']; qualColor=RGBColor(255,0,0)
        if muts.count(row[13])>1:
            comment=lang['Contamination']+'?'; commentColor=RGBColor(255,0,0)
            varsContam.append(str(i+rowNum))
    else:
        qual=lang['QualMod']; qualColor=RGBColor(255,78,0)
        varsForCheck.append(str(i+rowNum))
    if 'empty' in row[1]:
        varsForCheck.pop()
        comment=lang['Contamination']+'!'; commentColor=RGBColor(255,0,0)
    clinReasons=[]
    if 'Pathogenic' in row[27]: clinReasons.append('ClinVar')
    if 'yes' in row[25]: clinReasons.append('BIC')
    if len(clinReasons)==0: clinStatus='4'
    else: clinStatus='5'
    clinReasonsStr=' ,'.join(clinReasons)
    varPerc=str(int(round(float(row[18])*100,0)))+'%'
    reportRow=[str(i),row[1],row[10],row[13],row[14],qual,varPerc,clinStatus,clinReasonsStr,comment]
    cells=table.add_row().cells
    for j,r in enumerate(reportRow):
        cells[j].text=r
        if j==0 or j==5 or j==7:
            cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        set_cell_vertical_alignment(cells[j])
    cells[5].paragraphs[0].runs[0].font.color.rgb=qualColor
    cells[6].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    cells[8].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
    cells[9].paragraphs[0].runs[0].font.color.rgb=commentColor; cells[9].paragraphs[0].runs[0].font.bold=True
    cells[9].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
rowNum=i
# Predicted pathogenic
ws=wb.sheet_by_name('Predicted_pathogenic')
if ws.nrows>1:
    muts=[]
    for i in range(ws.nrows)[1:]:
        row=ws.row_values(i)
        muts.append(row[13])
    for i in range(ws.nrows)[1:]:
        if row[1] not in patVarNums.keys():
            patVarNums[row[1]]=1
        else:
            patVarNums[row[1]]+=1
        row=ws.row_values(i)
        comment=''; commentColor=RGBColor(0,0,0)
        if float(row[8])>=1200 and float(row[18])>=0.14: qual=lang['QualHigh']; qualColor=RGBColor(0,100,0)
        elif float(row[8])<500 and float(row[18])<0.14:
            varsForCheck.append(str(i+rowNum))
            qual=lang['QualLow']; qualColor=RGBColor(255,0,0)
            if muts.count(row[13])>1:
                comment=lang['Contamination']+'?'; commentColor=RGBColor(255,0,0)
                varsContam.append(str(i+rowNum))
        else:
            qual=lang['QualMod']; qualColor=RGBColor(255,78,0)
            varsForCheck.append(str(i+rowNum))
        if 'empty' in row[1]:
            varsForCheck.pop()
            comment=lang['Contamination']+'!'; commentColor=RGBColor(255,0,0)
        clinReasonsStr=lang['InSilicoTools']
        clinStatus='4'
        varPerc=str(int(round(float(row[18])*100,0)))+'%'
        reportRow=[str(rowNum+i),row[1],row[10],row[13],row[14],qual,varPerc,clinStatus,clinReasonsStr,comment]
        cells=table.add_row().cells
        for j,r in enumerate(reportRow):
            cells[j].text=r
            if j==0 or j==5 or j==7:
                cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            set_cell_vertical_alignment(cells[j])
        cells[5].paragraphs[0].runs[0].font.color.rgb=qualColor
        cells[6].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        cells[8].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
        cells[9].paragraphs[0].runs[0].font.color.rgb=commentColor; cells[9].paragraphs[0].runs[0].font.bold=True
        cells[9].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
rowNum=rowNum+i
# Unknown significance
ws=wb.sheet_by_name('Unknown')
if ws.nrows>1:
    muts=[]
    for i in range(ws.nrows)[1:]:
        row=ws.row_values(i)
        muts.append(row[13])
    for i in range(ws.nrows)[1:]:
##        if row[1] not in patVarNums.keys():
##            patVarNums[row[1]]=1
##        else:
##            patVarNums[row[1]]+=1
        row=ws.row_values(i)
        comment=''; commentColor=RGBColor(0,0,0)
        if float(row[8])>=1200 and float(row[18])>=0.14: qual=lang['QualHigh']; qualColor=RGBColor(0,100,0)
        elif float(row[8])<500 and float(row[18])<0.14:
            varsForCheck.append(str(i+rowNum))
            qual=lang['QualLow']; qualColor=RGBColor(255,0,0)
        else:
            qual=lang['QualMod']; qualColor=RGBColor(255,78,0)
            varsForCheck.append(str(i+rowNum))
        if 'empty' in row[1]:
            varsForCheck.pop()
            comment=lang['Contamination']+'!'; commentColor=RGBColor(255,0,0)
        clinReasonsStr='-'
        clinStatus='3'
        varPerc=str(int(round(float(row[18])*100,0)))+'%'
        reportRow=[str(rowNum+i),row[1],row[10],row[13],row[14],qual,varPerc,clinStatus,clinReasonsStr,comment]
        cells=table.add_row().cells
        for j,r in enumerate(reportRow):
            cells[j].text=r
            if j==0 or j==5 or j==7:
                cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            set_cell_vertical_alignment(cells[j])
        cells[5].paragraphs[0].runs[0].font.color.rgb=qualColor
        cells[6].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        cells[8].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
        cells[9].paragraphs[0].runs[0].font.color.rgb=commentColor; cells[9].paragraphs[0].runs[0].font.bold=True
        cells[9].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
# Definitions for Table
head4=doc.add_heading(lang['Designations'],2)
p=doc.add_paragraph(lang['CDS']+' – '+lang['CDS_Def'],style='List Bullet')
p=doc.add_paragraph(lang['Protein']+' – '+lang['Protein_Def'],style='List Bullet')
p=doc.add_paragraph(lang['VarPercent']+' – '+lang['VarPercent_Def'],style='List Bullet')
p=doc.add_paragraph(lang['ClinSign']+' – '+lang['ClinSign_Def'],style='List Bullet')
p=doc.add_paragraph(lang['ClassReasons']+' – '+lang['ClassReasons_Def'],style='List Bullet')
p=doc.add_paragraph(lang['Contamination']+' – '+lang['Contamination_Def'],style='List Bullet')
p=doc.add_paragraph('fs – frameshift',style='List Bullet')
# Warnings
head5=doc.add_heading(lang['Warnings'],2)
warningsColor=RGBColor(255,0,0)
head5.runs[0].font.color.rgb=warningsColor
if len(varsForCheck)>0:
    p=doc.add_paragraph(lang['VarForCheckWarn1']+', '.join(varsForCheck)+lang['VarForCheckWarn2'],style='List Bullet')
    p.runs[0].font.color.rgb=warningsColor
if len(varsContam)>0:
    p=doc.add_paragraph(lang['VarForCheckWarn1']+', '.join(varsContam)+lang['VarsContamWarn2'],style='List Bullet')
    p.runs[0].font.color.rgb=warningsColor
patsWithSeveralMuts=[]
for key,value in patVarNums.items():
    if value>1: patsWithSeveralMuts.append(key)
if len(patsWithSeveralMuts)>0:
    p=doc.add_paragraph(lang['PatsWithSeveralMutsWarn1']+', '.join(patsWithSeveralMuts)+lang['PatsWithSeveralMutsWarn2'],style='List Bullet')
    p.runs[0].font.color.rgb=warningsColor
changeStyles(doc)
doc.save(args.outFile)

